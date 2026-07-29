#!/usr/bin/env python3
"""锁妖塔 — 每日综合扫描 (1小时+日线双周期过滤)

流程:
  1. OKX取前N个活跃币种
  2. 1小时分析 → 88因子×99交易员 → 过滤(R1≥1.5%, OI≥600K)
  3. 日线分析 → 方向判断
  4. 综合过滤: 双周期一致 + ADX≥25 + TP1利润≥100%
  5. 输出标准表格 + 保存daily_picks.md

用法:
    python quant_factors/run_daily_picks.py
    python quant_factors/run_daily_picks.py --top 50 --min-r1 1.5
"""
import sys, os, json, urllib.request, time
import pandas as pd
import numpy as np
from datetime import datetime, timezone
from io import StringIO
from collections import OrderedDict

QF = os.path.dirname(os.path.abspath(__file__))
BASE = os.path.dirname(QF)
sys.path.insert(0, QF)
sys.path.insert(0, BASE)
from local_config import OKX_API_KEY
from okx_data_adapter import build_features_single
from capabilities import CAP_REGISTRY, evaluate_all

STABLE = {'USDT','USDC','DAI','TUSD','BUSD','FDUSD','USDP',
    'EUR','GBP','AUD','SGD','AED','CNY','JPY','KRW','USDG',
    'TRY','BRL','CAD','CHF','HKD','MXN','PI'}

# ═══════════════════════════════════════
# OKX API
# ═══════════════════════════════════════

def api_get(path):
    url = 'https://www.okx.com' + path
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    return json.loads(urllib.request.urlopen(req, timeout=15).read())


def fetch_list(top_n=15):
    r = api_get('/api/v5/market/tickers?instType=SPOT')
    coins = []
    for t in r.get('data', []):
        inst = t['instId']
        if not inst.endswith('-USDT'): continue
        base = inst.replace('-USDT', '')
        if base in STABLE: continue
        v = float(t.get('volCcy24h', '0') or 0)
        if v >= 500000:
            coins.append({'base': base, 'symbol': inst, 'vol': v})
    coins.sort(key=lambda x: x['vol'], reverse=True)
    return coins[:top_n]


def fetch_ohlc(symbol, bar='1H', limit=200):
    try:
        r = api_get(f'/api/v5/market/candles?instId={symbol}&bar={bar}&limit={limit}')
        raw = r.get('data', [])
        if not raw: return []
        raw.reverse()
        out = []
        for c in raw:
            ts = int(c[0]) / 1000
            dt = datetime.fromtimestamp(ts, tz=timezone.utc)
            out.append({'date': dt.strftime('%Y-%m-%d %H:%M'),
                        'open': float(c[1]), 'high': float(c[2]),
                        'low': float(c[3]), 'close': float(c[4]),
                        'volume': float(c[5])})
        return out
    except Exception as e:
        print(f'  [OKX] fetch_ohlc error: {e}')
        return []


def fetch_funding_rate(base):
    try:
        r = api_get(f'/api/v5/public/funding-rate?instId={base}-USDT-SWAP')
        if r.get('code') == '0' and r.get('data'):
            return float(r['data'][0]['fundingRate'])
        print(f'  [OKX] funding-rate for {base}: code={r.get("code")}')
    except Exception as e:
        print(f'  [OKX] funding-rate error for {base}: {e}')
    return 0.0


def fetch_open_interest(base):
    try:
        r = api_get(f'/api/v5/public/open-interest?instType=SWAP&instId={base}-USDT-SWAP')
        if r.get('code') == '0' and r.get('data'):
            return float(r['data'][0]['oi'])
        print(f'  [OKX] open-interest for {base}: code={r.get("code")}')
    except Exception as e:
        print(f'  [OKX] open-interest error for {base}: {e}')
    return 0.0


# ═══════════════════════════════════════
# 因子评分 (单行评估, 与run_5m_kol_consensus一致)
# ═══════════════════════════════════════

def cscore(cid, row, fr, oi):
    c = float(row['close']); h = float(row['high']); l = float(row['low']); o = float(row['open'])
    ma20 = float(row.get('ma20', c)); ma50 = float(row.get('ma50', c)); ma200 = float(row.get('ma200', c))
    rsi = float(row.get('rsi14', 50)); r1_ = float(row.get('r1', c)); s1_ = float(row.get('s1', c))
    fib = float(row.get('fib_618', c)); bbw = float(row.get('bb_width', 0))
    hist = float(row.get('macd_hist', 0)); pv = float(row.get('pivot', c))
    rng = h - l; body = abs(c - o)

    if cid == 'cap_044_regime_trending_up': return 0.6 if c > ma50 else -0.4
    if cid == 'cap_045_regime_trending_down': return 0.6 if c < ma50 else -0.4
    if cid == 'cap_046_regime_ranging': return 0.5 if rng > 0 and (h - l) / c < 0.005 else 0.0
    if cid == 'cap_047_regime_volatile': return 0.6 if rng > 0 and (h - l) / c > 0.01 else 0.0
    if cid == 'cap_070_parabolic_exhaustion': return 0.5 if abs(float(row.get('ret_5d', 0))) > 0.02 else 0.0
    if cid == 'cap_015_rsi_bullish_divergence': return 0.3 if rsi > 50 and c > ma50 else -0.1
    if cid == 'cap_016_rsi_bearish_divergence': return -0.3 if rsi < 50 and c < ma50 else 0.1
    if cid == 'cap_017_rsi_oversold_bounce': return 0.5 if rsi < 35 else (0.2 if rsi < 45 else 0)
    if cid == 'cap_018_ma_golden_cross': return 0.5 if ma50 > ma200 else -0.5
    if cid == 'cap_019_ma_death_cross': return -0.5 if ma50 < ma200 else 0.5
    if cid == 'cap_020_macd_histogram_cross': return 0.4 if hist > 0 else -0.4
    if cid == 'cap_021_bb_squeeze_breakout': return 0.4 if bbw < 0.003 else (-0.2 if bbw > 0.008 else 0.0)
    if cid == 'cap_022_fib_618_support':
        d = abs(c - fib) / c if c > 0 else 1; return 0.7 if d < 0.01 else (0.3 if d < 0.03 else 0.0)
    if cid == 'cap_069_moving_average_reclaim': return 0.6 if c > ma200 else -0.6
    if cid in ('emg_008_w50ema_bull_bear_divider', 'emg_014_horizontal_reclaim'): return 0.4 if c > ma50 else -0.4
    if cid == 'emg_022_200w_mechanical_buy': return 0.5 if c < ma200 * 0.9 else -0.2
    if cid == 'emg_029_200w_value_zone': return 0.5 if c < ma200 * 0.85 else -0.1
    if cid == 'emg_028_20w_200w_double_reclaim': return 0.4 if c > ma200 and c > ma50 else -0.3
    if cid == 'emg_001_quarterly_vwap': return 0.3 if c > pv else -0.3
    if cid in ('cap_037_halving_cycle', 'cap_038_4year_cycle', 'emg_005_4year_same_day_compare'): return -0.30
    if cid == 'emg_006_days_in_tight_range': return 0.1
    if cid == 'emg_023_monthly_seasonality': return 0.1 if datetime.now().month in (10, 11, 12, 1, 2, 3) else -0.1
    if cid in ('cap_023_elliott_wave_3', 'cap_024_wyckoff_accumulation_spring', 'cap_026_smc_order_block_retest'): return 0.3 if c > ma50 else -0.3
    if cid in ('cap_025_wyckoff_distribution_upthrust', 'cap_049_ict_fair_value_gap'): return -0.3 if c < ma50 else 0.3
    if cid in ('cap_001_falling_wedge_breakout', 'cap_003_bull_flag', 'cap_006_inverse_head_shoulders'): return 0.35 if c > ma50 else -0.2
    if cid in ('cap_002_rising_wedge_breakdown', 'cap_004_bear_flag', 'cap_005_head_shoulders_top'): return -0.35 if c < ma50 else 0.2
    if cid in ('cap_012_sfp', 'cap_014_trend_pullback', 'cap_052_liquidity_grab', 'cap_057_fake_breakout', 'emg_013_box_breakout'): return 0.3 if c > ma20 else -0.3
    if cid == 'cap_053_doji': return 0.3 if rng > 0 and body / rng < 0.1 else 0.0
    if cid == 'cap_054_engulfing': return 0.3 if c > o else 0.0
    if cid == 'cap_055_pin_bar':
        uw = h - max(c, o); lw = min(c, o) - l
        if rng > 0 and min(uw, lw) / rng < 0.1 and max(uw, lw) / rng > 0.5:
            return 0.4 if lw > uw else -0.4
        return 0.0
    if cid == 'cap_056_double_needle_bottom': return 0.3 if rsi < 40 else 0.0
    if cid in ('cap_027_dxy_inverse_btc', 'cap_028_spx_risk_on', 'cap_040_etf_flows_proxy'): return 0.1 if c > ma50 else -0.1
    if cid == 'cap_041_dont_catch_falling_knives': return -0.5 if rsi < 30 else 0.0
    if cid == 'cap_043_cut_losses_early': return -0.3 if c < s1_ else 0.0
    if cid == 'emg_009_range_middle_filter':
        pos = (c - s1_) / max(r1_ - s1_, 0.01); return 0.0 if 0.3 < pos < 0.7 else (0.3 if pos < 0.3 else -0.3)
    if cid == 'emg_031_keltner_mean_revert':
        kcu = float(row.get('kc_upper', c)); kcl = float(row.get('kc_lower', c)); adx14 = float(row.get('adx14', 25))
        if kcu - kcl <= 0: return 0.0
        dist_l = (c - kcl) / (kcu - kcl)
        if adx14 < 25 and dist_l > 0.8: return 0.5
        if adx14 < 25 and dist_l < 0.2: return -0.5
        return 0.0
    if cid == 'cap_031_funding_extreme_neg':
        if fr < -0.001: return 0.8
        if fr < -0.0005: return 0.6
        if fr < -0.0001: return 0.3
        return 0.0
    if cid == 'cap_032_funding_extreme_pos':
        if fr > 0.001: return -0.8
        if fr > 0.0005: return -0.6
        if fr > 0.0001: return -0.3
        return 0.0
    if cid == 'cap_033_oi_climb':
        if oi > 50000000: return 0.4
        if oi > 10000000: return 0.25
        if oi > 1000000: return 0.1
        return 0.0
    if cid == 'cap_059_funding_divergence':
        if fr > 0.0005 and rsi < 40: return 0.5
        if fr < -0.0005 and rsi > 60: return -0.5
        return 0.0
    return 0.0


def match_cap(pid, rids):
    if pid in rids: return pid
    parts = pid.split('_')
    if len(parts) >= 2 and parts[0] == 'cap' and parts[1].isdigit():
        m = [c for c in rids if c.startswith('cap_' + parts[1] + '_')]
        return m[0] if m else None
    return None


# ═══════════════════════════════════════
# KOL 投票引擎
# ═══════════════════════════════════════

def kol_vote(latest_row, reg, rids, profs, fr, oi, factor_scores=None, kol_weights=None):
    """全交易员投票, 返回 (long_n, short_n, neutral_n, avg_bias)
    
    Args:
        factor_scores: 可选, {cap_id: Series} 来自 evaluate_all() 的预计算分数。
                       优先使用, 找不到或 NaN 时回退到 cscore()。
        kol_weights:   可选, {handle: float} KOL权重, 默认1.0。
                       来自 kol_tracker.load_weights(), 持续预测错误的KOL自动降权。
    """
    fs = {}
    for cid in reg:
        try:
            if factor_scores is not None and cid in factor_scores:
                val = factor_scores[cid]
                # 兼容 Series(.iloc) 和 numpy array([-1])
                try:
                    if hasattr(val, 'iloc'):
                        v = float(val.iloc[-1])
                    elif hasattr(val, '__getitem__') and hasattr(val, '__len__'):
                        v = float(val[-1]) if len(val) > 0 else 0.0
                    else:
                        v = float(val)
                except Exception:
                    v = 0.0
                # ★ CAP_REGISTRY是事件型(仅触发日非零), cscore是状态型(始终有方向)
                # 当CAP_REGISTRY触发(非零)时用它, 否则回退cscore持续方向判断
                if not np.isnan(v) and v != 0:
                    fs[cid] = v
                    continue
            fs[cid] = cscore(cid, latest_row, fr, oi)
        except Exception as e:
            fs[cid] = 0.0

    tsigs = []
    for h, p in profs.items():
        tw, ws = 0.0, 0.0
        for cap in (p.get('capabilities_used', []) or []):
            rid = cap.get('id', ''); w = float(cap.get('weight', 0))
            mid = match_cap(rid, rids)
            if mid:
                s = fs.get(mid, 0)
                if s != 0:
                    ws += w * s
                    tw += abs(w)
        if tw > 0:
            sig = ws / tw
        else:
            b = p.get('bias_default', 'neutral')
            sig = 0.15 if b == 'long_tilted' else (-0.15 if b == 'short_tilted' else 0.0)
        tsigs.append(sig)

    # ★ 按KOL权重调整信号
    if kol_weights:
        w_arr = np.array([kol_weights.get(h, 1.0) for h in profs])
        tsigs_arr = np.array(tsigs)
        weighted_sigs = tsigs_arr * w_arr
        # 重新计算方向计数时使用加权信号
        arr = weighted_sigs
    else:
        arr = np.array(tsigs)

    ln = int(np.sum(arr > 0.03))
    sn = int(np.sum(arr < -0.03))
    nn = len(arr) - ln - sn
    avg = float(np.mean(arr))
    return ln, sn, nn, avg


# ═══════════════════════════════════════
# 单币分析 (1小时 + 日线)
# ═══════════════════════════════════════

def analyze_coin(base, reg, rids, profs, min_r1=1.5, min_oi=600000, lev_map=None, kol_weights=None):
    """对一个币做三重时间框架分析(15m+1H+日线), 返回结果dict或None"""
    sym = f'{base}-USDT'

    # ── LTF: 15m K线 (入场时机+短期KOL) ──
    cdl_15m = fetch_ohlc(sym, '15m', 200)
    if len(cdl_15m) < 20: return None

    df_15m = pd.DataFrame(cdl_15m)
    df_15m['date'] = pd.to_datetime(df_15m['date'])
    df_15m = df_15m.set_index('date').sort_index()
    feats_15m = build_features_single(df_15m)
    lat_15m = feats_15m.iloc[-1]
    cur = float(lat_15m['close'])

    # 衍生品
    fr = fetch_funding_rate(base)
    oi = fetch_open_interest(base)

    # 预计算因子分数 (CAP_REGISTRY 向量化评估)
    feats_15m = feats_15m.copy()
    feats_15m['funding_rate'] = fr
    feats_15m['open_interest'] = oi
    factor_scores_15m, _ = evaluate_all(feats_15m)

    # LTF KOL投票
    ltf_ln, ltf_sn, ltf_nn, ltf_avg = kol_vote(lat_15m, reg, rids, profs, fr, oi, factor_scores_15m, kol_weights)

    # ── MTF: 1H K线 (中期趋势仲裁, 新增) ──
    cdl_1h = fetch_ohlc(sym, '1H', 168)
    mtf_avg = 0.0; mtf_ln = mtf_sn = 0
    if len(cdl_1h) >= 20:
        df_1h = pd.DataFrame(cdl_1h)
        df_1h['date'] = pd.to_datetime(df_1h['date'])
        df_1h = df_1h.set_index('date').sort_index()
        feats_1h = build_features_single(df_1h)
        lat_1h = feats_1h.iloc[-1]

        # 注入衍生品列 + 预计算因子
        feats_1h = feats_1h.copy()
        feats_1h['funding_rate'] = fr
        feats_1h['open_interest'] = oi
        factor_scores_1h, _ = evaluate_all(feats_1h)

        mtf_ln, mtf_sn, _, mtf_avg = kol_vote(lat_1h, reg, rids, profs, fr, oi, factor_scores_1h, kol_weights)
        time.sleep(0.1)

    # ── HTF: 日线 (大趋势方向+Pivot) ──
    cdl_daily = fetch_ohlc(sym, '1D', 200)
    htf_avg = 0.0; htf_ln = htf_sn = 0
    if len(cdl_daily) >= 20:
        df_d = pd.DataFrame(cdl_daily)
        df_d['date'] = pd.to_datetime(df_d['date'])
        df_d = df_d.set_index('date').sort_index()
        feats_d = build_features_single(df_d)
        lat_d = feats_d.iloc[-1]

        # 注入衍生品列 + 预计算因子
        feats_d = feats_d.copy()
        feats_d['funding_rate'] = fr
        feats_d['open_interest'] = oi
        factor_scores_d, _ = evaluate_all(feats_d)

        htf_ln, htf_sn, _, htf_avg = kol_vote(lat_d, reg, rids, profs, fr, oi, factor_scores_d, kol_weights)
        time.sleep(0.15)

    # Pivot — 优先用日线范围, 回退到15m范围
    if len(cdl_daily) >= 20:
        hh = max(c['high'] for c in cdl_daily)
        ll = min(c['low'] for c in cdl_daily)
    else:
        hh = float(feats_15m['high'].max()); ll = float(feats_15m['low'].min())
    pv = (hh + ll + cur) / 3; r1 = 2 * pv - ll; r2 = pv + (hh - ll)
    s1 = 2 * pv - hh; s2 = pv - (hh - ll)
    r1_up = (r1 / cur - 1) * 100
    s2_down = (1 - s2 / cur) * 100

    # 杠杆
    okx_lev = lev_map.get(base, 20) if lev_map else 20

    # ADX + 趋势方向 (基于15m)
    adx = 0; adx_trend = ''; adx_hours_left = 0
    try:
        closes = np.array([c['close'] for c in cdl_15m])
        highs = np.array([c['high'] for c in cdl_15m])
        lows = np.array([c['low'] for c in cdl_15m])
        from smc_entry_signal import calc_adx
        adx, adx_rising = calc_adx(closes, highs, lows, 14)
        if adx > 0:
            ema7 = np.mean(closes[-7:])
            ema50 = np.mean(closes[-50:]) if len(closes) >= 50 else np.mean(closes)
            adx_trend = '↑' if (adx_rising and ema7 > ema50) else ('↓' if adx_rising else '→')
            if adx >= 25 and not adx_rising:
                adx_hours_left = max(0, (adx - 25) / adx * 4) if adx > 0 else 0
    except Exception as e:
        print(f'  [ADX] error for {base}: {e}')

    # 过滤1: R1≥min_r1, OI≥min_oi
    if r1_up < min_r1: return None
    if oi < min_oi: return None

    # 三重时间框架方向
    ltf_bias = 'long' if ltf_avg > 0.01 else ('short' if ltf_avg < -0.01 else 'neutral')
    mtf_bias = 'long' if mtf_avg > 0.01 else ('short' if mtf_avg < -0.01 else 'neutral')
    htf_bias = 'long' if htf_avg > 0.01 else ('short' if htf_avg < -0.01 else 'neutral')

    # 三重对齐度评分
    biases = [b for b in [ltf_bias, mtf_bias, htf_bias] if b != 'neutral']
    if len(biases) >= 2 and all(b == biases[0] for b in biases):
        alignment = 1.0       # 全一致 → 最高分
        alignment_grade = '三重一致'
    elif mtf_bias == htf_bias != 'neutral':
        alignment = 0.8       # MTF+HTF一致 → LTF分歧不重要
        alignment_grade = 'MTF+HTF一致'
    elif ltf_bias == mtf_bias != 'neutral':
        alignment = 0.6       # LTF+MTF一致 → HTF滞后
        alignment_grade = 'LTF+MTF一致'
    elif ltf_bias == htf_bias != 'neutral':
        alignment = 0.3       # LTF+HTF一致但MTF反 → 信心低
        alignment_grade = 'LTF+HTF一致(MTF分歧)'
    else:
        alignment = 0.0
        alignment_grade = '三向分歧'

    # 方向判定: MTF优先, 回退到HTF
    if mtf_bias != 'neutral':
        direction = mtf_bias
    elif htf_bias != 'neutral':
        direction = htf_bias
    elif ltf_bias != 'neutral':
        direction = ltf_bias
    else:
        return None

    # ★ 20维位置共识 (FMZ多策略融合, 百分制)
    try:
        from position_gauges import evaluate_all_positions
        pos_result = evaluate_all_positions(feats_15m, direction)
        pos_score = pos_result['score']       # 裁尾均值 0-100
        pos_lean = pos_result['lean']         # 偏向计数 (+N偏高/-N偏低)
        pos_grade = pos_result['grade']       # A/B/C/D/F
        pos_high = pos_result['high_count']
        pos_low = pos_result['low_count']
        pos_bias = pos_result['bias_mean']
    except Exception:
        pos_score = 50; pos_lean = 0; pos_grade = 'C'
        pos_high = 0; pos_low = 0; pos_bias = 0.0

    # ★ KOL共识强度 (按持仓4-8h加权: 1H主角/15m确认/日线方向)
    def _polarity(ln, sn):
        total = ln + sn
        return abs(ln - sn) / total if total > 0 else 0.0
    
    ltf_pol = _polarity(ltf_ln, ltf_sn)
    mtf_pol = _polarity(mtf_ln, mtf_sn)
    htf_pol = _polarity(htf_ln, htf_sn)
    
    kol_consensus = mtf_pol * 0.50 + ltf_pol * 0.30 + htf_pol * 0.20
    
    # 方向矛盾惩罚: 15m反对1H→×0.7, 日线反对1H→×0.5
    if ltf_bias != 'neutral' and mtf_bias != 'neutral' and ltf_bias != mtf_bias:
        kol_consensus *= 0.7
    if htf_bias != 'neutral' and mtf_bias != 'neutral' and htf_bias != mtf_bias:
        kol_consensus *= 0.5

    # TP1利润
    if direction == 'long':
        tp1_pct = r1_up
    else:
        tp1_pct = s2_down

    tp1_profit = tp1_pct * okx_lev

    # TP1可行性
    tp1_req_pct = 100.0 / okx_lev  # 翻倍所需%
    liq_ratio = cur / abs(cur - s2) if abs(cur - s2) > 0.001 else 1
    safe_lev = (okx_lev <= liq_ratio)
    atr_val = float(lat_15m.get('atr14', 0))
    if cur > 0 and atr_val > 0:
        atr_daily_pct = (atr_val / cur * 100) * 96 * 0.7
        days_to_tp1 = tp1_req_pct / atr_daily_pct if atr_daily_pct > 0 else 999
    else:
        days_to_tp1 = 999
    tp1_score = 0
    if safe_lev: tp1_score += 4
    tp1_score += max(0, min(3, int(3 - (days_to_tp1 - 1))))
    if adx >= 25: tp1_score += 3
    elif adx >= 12: tp1_score += 1

    # 方向一致性 (已由三重框架覆盖)
    consistent = alignment >= 0.6

    # 市场状态
    if adx >= 25: state = '趋势'
    elif adx >= 12: state = '过渡'
    else: state = '震荡'

    # ORB区间（最近12根15分钟K线）
    orb_low = min(c['low'] for c in cdl_15m[-13:-1]) if len(cdl_15m) >= 14 else cur * 0.99
    orb_high = max(c['high'] for c in cdl_15m[-13:-1]) if len(cdl_15m) >= 14 else cur * 1.01

    return {
        'base': base,
        'entry': cur,
        'r1': r1, 'r2': r2, 's1': s1, 's2': s2,
        'r1_up': r1_up, 's2_down': s2_down,
        'orb_low': orb_low, 'orb_high': orb_high,
        # 三重时间框架KOL
        'ltf_avg': ltf_avg, 'ltf_bias': ltf_bias,
        'ltf_long': ltf_ln, 'ltf_short': ltf_sn, 'ltf_neutral': ltf_nn,
        'mtf_avg': mtf_avg, 'mtf_bias': mtf_bias,
        'mtf_long': mtf_ln, 'mtf_short': mtf_sn,
        'htf_avg': htf_avg, 'htf_bias': htf_bias,
        'htf_long': htf_ln, 'htf_short': htf_sn,
        # 对齐度
        'alignment': alignment, 'alignment_grade': alignment_grade,
        # 向后兼容 (旧字段名, 供未更新代码使用)
        'm5_avg': ltf_avg, 'm5_bias': ltf_bias,
        'm5_long': ltf_ln, 'm5_short': ltf_sn, 'm5_neutral': ltf_nn,
        'daily_avg': htf_avg, 'daily_bias': htf_bias,
        'daily_long': htf_ln, 'daily_short': htf_sn,
        'consistent': consistent,
        'adx': adx, 'adx_trend': adx_trend, 'adx_hours_left': adx_hours_left,
        'rsi': float(lat_15m.get('rsi14', 50)),
        # ★ 20维位置共识 (百分制, 替代单维度kc_pos)
        'kc_pos': pos_score / 100.0,       # 兼容旧代码: 转为[0,1]
        'pos_score': pos_score,            # 裁尾均值 0-100
        'pos_lean': pos_lean,              # 偏向 +N偏高/-N偏低
        'pos_grade': pos_grade,            # A/B/C/D/F
        'pos_high_count': pos_high,
        'pos_low_count': pos_low,
        'pos_bias': pos_bias,
        'kol_consensus': round(kol_consensus, 4),   # KOL共识强度(持仓周期加权)
        'funding_rate': fr, 'open_interest': oi,
        'okx_lev': okx_lev,
        'tp1_pct': tp1_pct,
        'tp1_profit': tp1_profit,
        'tp1_score': tp1_score,
        'tp1_pass': tp1_score >= 5,
        'market_state': state,
        'direction': direction,
        'fr_pct': fr * 100,
        'df_15m': cdl_15m,  # 原始15m K线数据，供审判系统复用
    }


# ═══════════════════════════════════════
# 格式化
# ═══════════════════════════════════════

def fmt_price(p):
    if p > 10: return f'${p:.2f}'
    elif p > 1: return f'${p:.4f}'
    elif p > 0.001: return f'${p:.6f}'
    else: return f'${p:.8f}'


def fmt_pos(kc_pos, pos_high=0, pos_low=0, pos_lean=0):
    """20维位置共识 → emoji + 百分制 + 偏向"""
    score = int(round(kc_pos * 100)) if kc_pos <= 1.0 else int(kc_pos)
    
    # 偏向字符串
    if pos_lean > 0:
        lean_str = f'+{pos_lean}'
    elif pos_lean < 0:
        lean_str = f'{pos_lean}'
    else:
        lean_str = '±0'
    
    if score >= 75:
        return f'🔴极高{score}({lean_str})'
    elif score >= 60:
        return f'🟡偏高{score}({lean_str})'
    elif score > 40:
        return f'🟢中位{score}({lean_str})'
    elif score >= 25:
        return f'🔵偏低{score}({lean_str})'
    else:
        return f'🟣极低{score}({lean_str})'


def fmt_target(r, is_long):
    """取TP1目标价"""
    if is_long:
        return r['r1']
    else:
        return r['s2']


def run(top_n=50, min_r1=1.5, min_oi=600000, coins=None):
    print()
    print('  ╔══════════════════════════════════════╗')
    print('  ║   锁妖塔 — 每日综合扫描              ║')
    now_str_short = datetime.now().strftime('%Y-%m-%d %H:%M')
    print(f'  ║   {now_str_short}                 ║')
    print('  ╚══════════════════════════════════════╝')
    print(f'  因子: {len(CAP_REGISTRY)} | 交易员: 99')

    # ★ 加载KOL绩效权重
    try:
        from kol_tracker import load_weights, update_after_backtest
        kol_weights = load_weights(profs)
        kw_active = sum(1 for w in kol_weights.values() if w != 1.0)
        if kw_active > 0:
            print(f'  KOL权重: {kw_active}个已校准')
    except Exception:
        kol_weights = None

    # 将锁妖塔扫描详情存入缓冲区，等重要的结论先出
    _buf = StringIO()
    _old_stdout = sys.stdout
    sys.stdout = _buf

    print(f'  扫描: 前{top_n}币 | 过滤: R1≥{min_r1}% OI≥{min_oi/1e6:.1f}M')


    reg = CAP_REGISTRY; rids = set(reg.keys())

    # 加载交易员
    profs = {}
    pd_ = os.path.join(BASE, 'profiles_v2')
    for f in sorted(os.listdir(pd_)):
        if f.endswith('.json'):
            try:
                p = json.load(open(os.path.join(pd_, f), encoding='utf-8'))
                profs[f.replace('.json', '')] = p
            except Exception as e:
                print(f'  [profiles] load {f} error: {e}')

    if coins:
        coins_list = [{'base': c.upper(), 'symbol': c.upper() + '-USDT', 'vol': 0} for c in coins]
        print(f'  指定币种: {len(coins_list)}个')
    else:
        coins_list = fetch_list(top_n)
        print(f'  币种: {len(coins_list)}')
    print()

    # 预取杠杆数据
    lev_map = {}
    try:
        inst = api_get('/api/v5/public/instruments?instType=SWAP')
        if inst.get('code') == '0':
            for d in inst.get('data', []):
                di = d['instId']
                if di.endswith('-USDT-SWAP'):
                    base_n = di.replace('-USDT-SWAP', '')
                    lev_map[base_n] = int(d.get('lever', 20))
    except Exception as e:
        print(f'  [OKX] instrument/leverage error: {e}')
    print(f'  杠杆数据: {len(lev_map)}个币')

    results = []
    for i, coin in enumerate(coins_list):
        base = coin['base']
        r = analyze_coin(base, reg, rids, profs, min_r1, min_oi, lev_map, kol_weights)
        if r:
            results.append(r)
            arrow = '✅' if r['consistent'] else '⚠️'
            align_tag = r.get('alignment_grade', '')[:6]
            print(f'  [{i+1}/{len(coins_list)}] {base:<8s} 15m L/S={r["m5_long"]}/{r["m5_short"]:<2d} '
                  f'1H L/S={r["mtf_long"]}/{r["mtf_short"]:<2d} '
                  f'D L/S={r["daily_long"]}/{r["daily_short"]:<2d} '
                  f'ADX={r["adx"]:.0f} R1={r["r1_up"]:.1f}% Pft={r["tp1_profit"]:.0f}% '
                  f'{arrow} {align_tag}')
        else:
            print(f'  [{i+1}/{len(coins_list)}] {base:<8s} 跳过')
        time.sleep(0.1)

    # ── 综合过滤 ──
    passed = [r for r in results if r['consistent'] and r['adx'] >= 25 and r['tp1_profit'] >= 100]
    short_ok = [r for r in passed if r['direction'] == 'short']
    long_ok = [r for r in passed if r['direction'] == 'long']

    # ── 大盘环境过滤 ──
    consistent_results = [r for r in results if r['consistent']]
    long_count = sum(1 for r in consistent_results if r['direction'] == 'long')
    short_count = sum(1 for r in consistent_results if r['direction'] == 'short')
    if short_count >= long_count * 2 and short_count >= 3:
        long_ok = []  # 市场偏空，不做多
        market_env = '🔴 市场偏空'
    elif long_count >= short_count * 2 and long_count >= 3:
        short_ok = []  # 市场偏多，不做空
        market_env = '🟢 市场偏多'
    else:
        market_env = '⚪ 市场均衡'

    short_ok.sort(key=lambda r: r['tp1_profit'], reverse=True)
    long_ok.sort(key=lambda r: r['tp1_profit'], reverse=True)

    # ── 输出 ──
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    sep = '=' * 90
    dash2 = '-' * 95

    # 排名：按三重对齐度+方向分组
    ranking_groups = []
    # 全一致最优 → MTF+HTF一致次之 → 两两一致 → 三向分歧垫底
    for grade_order, grade_label in [(1.0, '三重一致'), (0.8, 'MTF+HTF一致'), (0.6, 'LTF+MTF一致'), (0.3, 'LTF+HTF一致'), (0.0, '三向分歧')]:
        group = [r for r in results if r.get('alignment', 0) == grade_order]
        if group:
            direction_order = {'long': 0, 'short': 1, 'neutral': 2}
            group.sort(key=lambda r: (direction_order.get(r['direction'], 2), -r['adx']))
            ranking_groups.extend(group)
    ranked = ranking_groups
    # 未设置alignment的旧数据回退到日线分组
    if not ranking_groups:
        long_group = [r for r in results if r['daily_avg'] > 0.01]
        short_group = [r for r in results if r['daily_avg'] < -0.01]
        neutral_group = [r for r in results if -0.01 <= r['daily_avg'] <= 0.01]
        long_group.sort(key=lambda r: -r['adx'])
        short_group.sort(key=lambda r: -r['adx'])
        neutral_group.sort(key=lambda r: -r['adx'])
        ranked = long_group + short_group + neutral_group

    print(f'\n\n  {sep}')
    print(f'  ★ 全量排名 (按信号 空←→多 排列)')
    print(f'  {now_str}')
    print(f'  {sep}')
    print(f'  扫描: {len(coins_list)}币 → 基础通过: {len(results)}')
    print(f'  ✅=方向一致  ⚠️=方向冲突')
    print()
    hdr = (f'  {"序":>2s} {"币种":<6s} {"方向":>2s} {"强度":>5s} {"方向":>2s} {"小时":>4s} {"位置":<6s}'
           f' {"RSI":>4s} {"入场价":>10s} {"止盈价":>10s} {"杠杆":>4s} {"TP1利润%":>8s}'
           f' {"15mKOL":>7s} {"1HKOL":>5s} {"日KOL":>6s} {"对齐":>6s}')
    print(hdr)
    print(f'  {dash2}')
    for i, r in enumerate(ranked, 1):
        arrow_m = '✅' if r['consistent'] else '⚠️'
        align_short = r.get('alignment_grade', '')[:6]
        hrs_str = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
        pos_str = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
        if r['direction'] == 'long':
            target = r['r1']
            d_arrow = '🟢'
        else:
            target = r['s2']
            d_arrow = '🔴'
        profit_str = f'{r["tp1_profit"]:.0f}%'
        print(f'  {i:>2d} {r["base"]:<6s} {d_arrow:>2s} {r["adx"]:>5.1f} {r.get("adx_trend",""):>2s} {hrs_str:>4s} {pos_str:<6s}'
              f' {r["rsi"]:>4.0f} {fmt_price(r["entry"]):>10s} {fmt_price(target):>10s}'
              f' {r["okx_lev"]:>3d}x {profit_str:>8s}'
              f' {r["m5_long"]}/{r["m5_short"]:>3d} {r["mtf_long"]}/{r["mtf_short"]:>2d}'
              f' {r["daily_long"]}/{r["daily_short"]:>3d}'
              f' {arrow_m:>2s}{align_short:>4s}')
    print()

    # ── 推荐 (严格过滤 TOP3) ──
    print(f'  {sep}')
    print(f'  ★ 今日推荐 {market_env}')
    print(f'  {now_str}')
    print(f'  {sep}')
    print()

    if short_ok:
        print(f'  ── 做空推荐 ──')
        hdr = (f'  {"序":>2s} {"币种":<6s} {"强度":>5s} {"方向":>2s} {"小时":>4s} {"位置":<6s}'
               f' {"RSI":>4s} {"入场价":>10s} {"止盈价":>10s} {"杠杆":>4s} {"TP1利润%":>8s}'
               f' {"15mKOL":>7s} {"1HKOL":>5s} {"日KOL":>6s} {"对齐":>4s}')
        print(hdr)
        print(f'  {dash2}')
        for i, r in enumerate(short_ok[:5], 1):
            target = fmt_target(r, False)
            profit_str = f'{r["tp1_profit"]:.0f}%'
            hrs_str = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
            pos_str = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
            align_str = r.get('alignment_grade', '')[:4]
            print(f'  {i:>2d} {r["base"]:<6s} {r["adx"]:>5.1f} {r.get("adx_trend",""):>2s} {hrs_str:>4s} {pos_str:<6s}'
                  f' {r["rsi"]:>4.0f} {fmt_price(r["entry"]):>10s} {fmt_price(target):>10s}'
                  f' {r["okx_lev"]:>3d}x {profit_str:>8s}'
                  f' {r["m5_long"]}/{r["m5_short"]:>3d} {r["mtf_long"]}/{r["mtf_short"]:>2d}'
                  f' {r["daily_long"]}/{r["daily_short"]:>3d} {align_str:>4s}')
        print()

    if long_ok:
        print(f'  ── 做多推荐 ──')
        hdr = (f'  {"序":>2s} {"币种":<6s} {"强度":>5s} {"方向":>2s} {"小时":>4s} {"位置":<6s}'
               f' {"RSI":>4s} {"入场价":>10s} {"止盈价":>10s} {"杠杆":>4s} {"TP1利润%":>8s}'
               f' {"15mKOL":>7s} {"1HKOL":>5s} {"日KOL":>6s} {"对齐":>4s}')
        print(hdr)
        print(f'  {dash2}')
        for i, r in enumerate(long_ok[:5], 1):
            target = fmt_target(r, True)
            profit_str = f'{r["tp1_profit"]:.0f}%'
            hrs_str = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
            pos_str = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
            print(f'  {i:>2d} {r["base"]:<6s} {r["adx"]:>5.1f} {r.get("adx_trend",""):>2s} {hrs_str:>4s} {pos_str:<6s}'
                  f' {r["rsi"]:>4.0f} {fmt_price(r["entry"]):>10s} {fmt_price(target):>10s}'
                  f' {r["okx_lev"]:>3d}x {profit_str:>8s}'
                  f' {r["m5_long"]}/{r["m5_short"]:>3d} {r["mtf_long"]}/{r["mtf_short"]:>2d}'
                  f' {r["daily_long"]}/{r["daily_short"]:>3d} {r.get("alignment_grade","")[:4]:>4s}')
        print()

    # ── 分批入场参考 ──
    if short_ok or long_ok:
        print(f'  ── 分批入场参考 ──')
        for r in (short_ok[:3] + long_ok[:5])[:5]:
            is_long = r['direction'] == 'long'
            dir_cn = '做多' if is_long else '做空'
            liq = r['entry'] - (r['entry'] / r['okx_lev']) if is_long else r['entry'] + (r['entry'] / r['okx_lev'])
            stop = liq + r['entry'] * 0.002 * (1 if is_long else -1)
            orb_ref = r.get('orb_low', r['entry'] * 0.99) if is_long else r.get('orb_high', r['entry'] * 1.01)
            if is_long:
                orb_ref = max(orb_ref, stop + r['entry'] * 0.001)
            else:
                orb_ref = min(orb_ref, stop - r['entry'] * 0.001)
            tag = 'ORB下沿' if is_long else 'ORB上沿'
            target = r['r1'] if is_long else r['s2']
            hrs = r.get('adx_hours_left', 0)
            align = r.get('alignment_grade', '')
            hrs_s = f' | 趋势 {hrs:.0f}h' if hrs > 0 else ''
            align_s = f' | {align}' if align else ''
            print(f'  ▶ {r["base"]} {dir_cn}{align_s}')
            print(f'    入场1 {fmt_price(r["entry"])}(20U) → 入场2 {fmt_price(orb_ref)}(30U) {tag}{hrs_s}')
            print(f'    止损 {fmt_price(stop)} | TP1 {fmt_price(target)}')
        print()

    # ── 未通过的原因统计 ──
    total = len(results)
    no_consistency = sum(1 for r in results if not r['consistent'])
    no_adx = sum(1 for r in results if r['consistent'] and r['adx'] < 25)
    no_profit = sum(1 for r in results if r['consistent'] and r['adx'] >= 25 and r['tp1_profit'] < 100)

    print(f'  {dash2}')
    print(f'  过滤统计:')
    print(f'    - 基础通过(R1≥{min_r1}%+OI≥{min_oi/1e6:.1f}M): {total}')
    print(f'    - 方向冲突排除: {no_consistency}')
    print(f'    - ADX<25排除: {no_adx}')
    print(f'    - TP1利润<100%排除: {no_profit}')
    print(f'    - 最终推荐: {len(passed)}')
    print()

    # 恢复 stdout，保存扫描输出到变量
    sys.stdout = _old_stdout
    _scan_output = _buf.getvalue()

    # ── 综合推荐 + 精准入场（在审判之前，是最终结论）──
    try:
        from judge_system.run_judge import run_judge, ensure_detectors_registered
        ensure_detectors_registered()
        coins_data = {}
        for r in results:
            if 'df_15m' in r:
                from okx_data_adapter import normalize_ohlc_df
                raw = r['df_15m']
                df = normalize_ohlc_df(raw)
                coins_data[r['base']] = df
        coin_symbols = [c['base'] for c in coins_list] if coins_list else None
        verdicts, dis_count = run_judge(top_n=top_n, coins=coin_symbols, compare=True, table_only=True,
                                        coins_data=coins_data, print_verdict=False,
                                        bars_per_unit=96)  # ★ 15m数据=96根/天

        # ── 综合推荐：锁妖塔 + 审判系统 ──
        if verdicts and passed:
            print(f'  {"=" * 60}')
            print(f'  ★ 综合推荐 Top 3 — 锁妖塔 + 审判系统融合')
            print(f'  {"=" * 60}')

            # 对每个通过的推荐，综合评分
            combined = []
            for r in passed:
                base = r['base']
                v = verdicts.get(base)

                # 锁妖塔评分（KOL共识+趋势+对齐+利润）
                kol_c = r.get('kol_consensus', 0.5)
                p_score = (min(r['adx'] / 50, 1.0) * 0.25 +    # 趋势强度
                           r.get('alignment', 0.5) * 0.25 +      # 三重对齐
                           kol_c * 0.30 +                        # KOL共识(持仓周期加权)
                           min(r['tp1_profit'] / 500, 1.0) * 0.20)  # 利润空间
                p_dir = r['direction']

                # 反转向量 + 趋势强度
                ts_score = 0.0
                ts_raw = 0.0
                df_coin = coins_data.get(base)
                if df_coin is not None:
                    from judge_system.entry_planner import check_trend_strength
                    ts = check_trend_strength(df_coin, p_dir)
                    ts_raw = ts.get('strength', 0)
                    # 软化: 逆势(<0)淘汰, 弱趋势(0~0.2)打折, 强趋势正常
                    if ts_raw <= 0:
                        ts_score = -999  # 标记淘汰
                    elif ts_raw <= 0.2:
                        ts_score = ts_raw / 4 * 0.5  # 弱趋势打五折
                    else:
                        ts_score = max(0, ts_raw / 4)

                if v and abs(v.judge_score) > 0.1:
                    j_dir = v.judge_direction
                    # ★ 审判直接用confidence, 不再打折
                    j_score = v.judge_confidence

                    if p_dir == j_dir:
                        if ts_score < 0:
                            final_score = -1.0
                            tag = '❌逆势'
                        else:
                            # 四维综合: 锁妖塔35% + 审判30% + 趋势20% + 位置15%
                            pos_score_val = 0.0
                            if direction == 'long':
                                pos_score_val = max(0, (50 - r.get('pos_score', 50)) / 25)
                            else:
                                pos_score_val = max(0, (r.get('pos_score', 50) - 50) / 25)
                            pos_score_val = min(1.0, pos_score_val)
                            final_score = (p_score * 0.35 + j_score * 0.30 +
                                           ts_score * 0.20 + pos_score_val * 0.15)
                            pos_grade = r.get('pos_grade', 'C')
                            if pos_grade in ('D', 'F'):
                                tag = f'⚠️一致(位置{pos_grade}级)'
                                final_score *= 0.3   # ★ 位置严重冲突 → 扣70%
                            elif pos_grade == 'C':
                                tag = '✅一致'
                            else:  # A or B
                                tag = '✅一致'
                                final_score *= 1.05  # 位置支持方向 → 小幅奖励
                    else:
                        # 方向相反 → 不推荐
                        final_score = -1.0
                        tag = '❌分歧'
                else:
                    # 审判系统无信号，只用锁妖塔
                    final_score = p_score * 0.5
                    tag = '⚪仅锁妖塔'

                combined.append({
                    'base': base,
                    'direction': p_dir,
                    'score': final_score,
                    'tag': tag,
                    'entry': r['entry'],
                    'tp1': r['r1'] if p_dir == 'long' else r['s2'],
                    'adx': r['adx'],
                    'kol': f'{r["m5_long"]}/{r["m5_short"]}',
                    'kol_1h': f'{r.get("mtf_long",0)}/{r.get("mtf_short",0)}',
                    'alignment': r.get('alignment_grade', ''),
                })

            # 过滤掉分歧的，按综合评分排序
            valid = [c for c in combined if c['score'] > 0]
            valid.sort(key=lambda x: x['score'], reverse=True)

            for i, c in enumerate(valid[:3], 1):
                dir_arrow = '🟢做多' if c['direction'] == 'long' else '🔴做空'
                align_str = c.get('alignment', '')
                print(f'  {i}. {c["base"]:<6} {dir_arrow}  '
                      f'评分={c["score"]:.2f}  ADX={c["adx"]:.0f}  '
                      f'入场={fmt_price(c["entry"])}  TP1={fmt_price(c["tp1"])}  '
                      f'KOL={c["kol"]} 1HKOL={c["kol_1h"]}  '
                      f'{align_str}  {c["tag"]}')

            disagreements = [c for c in combined if c['score'] < 0]
            if disagreements:
                print(f'  {"-" * 60}')
                print(f'  ⚠️ 分歧排除: {", ".join(c["base"] for c in disagreements)}')

            print(f'  {"=" * 60}')
            print()

            # ── 精准入场方案 + 收集回测数据 ──
            entry_plans = {}  # 用于回测保存
            for c in valid[:3]:
                base = c['base']
                r = next((x for x in results if x['base'] == base), None)
                if r and 'df_15m' in r:
                    from okx_data_adapter import normalize_ohlc_df
                    raw = r['df_15m']
                    df = normalize_ohlc_df(raw)

                    from judge_system.entry_planner import plan_entry, print_entry_plan
                    plan = plan_entry(base, c['direction'], c['entry'],
                                      r.get('okx_lev', 20), df)
                    print(f'  --- 精准入场: {base} ---')
                    print_entry_plan(plan)
                    entry_plans[base] = plan

            # ── 保存本次推荐到回测日志 ──
            try:
                from judge_system.backtest_log import save_recommendation
                rec_data = []
                for c in valid[:3]:
                    r = next((x for x in results if x['base'] == c['base']), None)
                    plan = entry_plans.get(c['base'])
                    # ★ 修复: TP1/liq 使用 entry_planner 的精准计算值, 不是 pivot
                    tp1_val = plan['tp1_price'] if plan else c['tp1']
                    liq_val = plan['liq_price'] if plan else (c['entry'] + c['entry']/r.get('okx_lev',20)*0.7 if c['direction']=='short'
                               else c['entry'] - c['entry']/r.get('okx_lev',20)*0.7)
                    rec_data.append({
                        'base': c['base'],
                        'direction': c['direction'],
                        'entry': c['entry'],
                        'tp1': tp1_val,
                        'liq': liq_val,
                        'lev': r.get('okx_lev', 20) if r else 20,
                        # ★ 新增: 记录推荐时的上下文
                        'adx': r.get('adx', 0) if r else 0,
                        'rsi': r.get('rsi', 0) if r else 0,
                        'kol_15m': f"{r.get('ltf_long',0)}/{r.get('ltf_short',0)}",
                        'kol_1h': f"{r.get('mtf_long',0)}/{r.get('mtf_short',0)}",
                        'kol_d': f"{r.get('htf_long',0)}/{r.get('htf_short',0)}",
                        'alignment': r.get('alignment_grade', ''),
                        # ★ 保存15m K线快照供逐K线回放
                        'df_15m_snapshot': r.get('df_15m', []) if r else [],
                    })
                save_recommendation(rec_data)
            except Exception as e:
                print(f'  [backtest] save error: {e}')

        # ── 回测上次推荐（最近5次）──
        try:
            from judge_system.backtest_log import init as bt_init, backtest_last, backtest_today, print_backtest
            bt_init(QF)
            now_hour = datetime.now().hour
            if 11 <= now_hour <= 13:
                bt_results, bt_summary, bt_metrics = backtest_today()
                if bt_results:
                    print_backtest(bt_results, bt_summary, bt_metrics)
            bt_results, bt_summary, bt_metrics = backtest_last()
            if bt_results:
                print_backtest(bt_results, bt_summary, bt_metrics)
                # ★ 用回测结果更新KOL绩效
                if kol_weights and bt_results:
                    try:
                        update_after_backtest(bt_results, profs)
                    except Exception:
                        pass
        except Exception as e:
            print(f'  [回测] 跳过: {e}')

        # ── 锁妖塔扫描详情（放在结论之后）──
        # 将在 finally 块中统一打印

        # ── 审判系统验证表（放在最后供参考）──
        if verdicts:
            from judge_system.run_judge import print_verdict_table
            from judge_system.run_judge import load_pagoda_results
            pagoda_results = load_pagoda_results()
            print_verdict_table(verdicts, pagoda_results)

    except Exception as e:
        print(f'  [审判] 跳过: {e}')
    finally:
        # 确保 stdout 恢复，扫描详情打印
        if '_old_stdout' in locals():
            sys.stdout = _old_stdout
        if '_scan_output' in locals():
            print(_scan_output)
    print()

    # ── 保存 ──
    op = os.path.join(QF, 'daily_picks.json')
    # 去掉 df_15m 字段（只用于审判系统，无需保存到JSON）
    passed_clean = []
    for r in passed:
        clean = {k: v for k, v in r.items() if k != 'df_15m'}
        passed_clean.append(clean)
    json.dump(passed_clean, open(op, 'w', encoding='utf-8'), ensure_ascii=False, indent=2)

    # 保存Markdown
    md_lines = [
        f'# 锁妖塔每日推荐 — {datetime.now().strftime("%Y-%m-%d %H:%M")}',
        f'',
        f'## 概览',
        f'- 扫描: {len(coins_list)}币 → 通过: {total} → 最终推荐: {len(passed)}',
        f'- 执行时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}',
        f'- 过滤标准: 双周期一致 + ADX≥25 + TP1利润≥100%',
        f'',
    ]
    if short_ok:
        md_lines.append('## 🔴 做空推荐')
        md_lines.append('| # | 币种 | ADX | RSI | 入场价 | 止盈价 | 涨幅% | 杠杆 | TP1利润 | KOL(15m) | KOL(日) | 费率 |')
        md_lines.append('|---|:----:|:---:|:---:|:------:|:------:|:-----:|:----:|:-------:|:-------:|:-------:|:----:|')
        for i, r in enumerate(short_ok[:5], 1):
            t = fmt_target(r, False)
            md_lines.append(f'| {i} | {r["base"]} | {r["adx"]:.1f} | {r["rsi"]:.0f} | {fmt_price(r["entry"])} | {fmt_price(t)} | {r["s2_down"]:+.2f}% | {r["okx_lev"]}x | {r["tp1_profit"]:.0f}% | {r["m5_long"]}/{r["m5_short"]} | {r["daily_long"]}/{r["daily_short"]} | {r["fr_pct"]:+.4f}% |')
    if long_ok:
        md_lines.append('')
        md_lines.append('## 🟢 做多推荐')
        md_lines.append('| # | 币种 | ADX | RSI | 入场价 | 止盈价 | 涨幅% | 杠杆 | TP1利润 | KOL(15m) | KOL(日) | 费率 |')
        md_lines.append('|---|:----:|:---:|:---:|:------:|:------:|:-----:|:----:|:-------:|:-------:|:-------:|:----:|')
        for i, r in enumerate(long_ok[:5], 1):
            t = fmt_target(r, True)
            md_lines.append(f'| {i} | {r["base"]} | {r["adx"]:.1f} | {r["rsi"]:.0f} | {fmt_price(r["entry"])} | {fmt_price(t)} | {r["r1_up"]:+.2f}% | {r["okx_lev"]}x | {r["tp1_profit"]:.0f}% | {r["m5_long"]}/{r["m5_short"]} | {r["daily_long"]}/{r["daily_short"]} | {r["fr_pct"]:+.4f}% |')
    md_path = os.path.join(QF, 'daily_picks.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    # 保存Word文档
    try:
        save_to_docx(passed, short_ok, long_ok, results, len(coins_list), len(results), now_str, QF, market_env)
    except Exception as e:
        print(f'  Word生成跳过: {e}')
    print()

    return passed


def save_to_docx(passed, short_ok, long_ok, all_results, total_coins, total_passed, now_str, QF, market_env=''):
    """生成Word交易报告 (横板)"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    doc = Document()
    
    # 横板
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── 区块1：报告头 ──
    title = doc.add_heading(f'锁妖塔每日交易报告 {market_env}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'生成时间: {now_str}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 概览
    info = doc.add_paragraph()
    info.add_run(f'扫描币种: {total_coins}个').bold = True
    info.add_run(f'  →  基础通过: {total_passed}个')
    noc = sum(1 for r in passed if not r.get('consistent', True))
    noa = sum(1 for r in passed if r.get('consistent', True) and r.get('adx',0) < 25)
    nop = sum(1 for r in passed if r.get('consistent', True) and r.get('adx',0) >= 25 and r.get('tp1_profit',0) < 100)
    info.add_run(f'  →  最终推荐: {len(passed)}个')

    doc.add_paragraph()

    # ── 通用表格样式 ──
    def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=8):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.bold = bold
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    def add_heading_row(table, headers):
        row = table.rows[0]
        for i, h in enumerate(headers):
            set_cell_text(row.cells[i], h, bold=True, size=9)

    def add_data_row(table, row_idx, values):
        row = table.rows[row_idx]
        for i, v in enumerate(values):
            set_cell_text(row.cells[i], str(v), size=8)

    # ── 区块2：全量排名TOP10 ──
    if all_results:
        sorted_results = sorted(all_results, key=lambda r: -r.get('alignment', 0))
        doc.add_heading('全量排名 (对齐度↓)', level=1)
        headers = ['序','币种','方向','强度','方向','小时','位置','RSI','入场价','止盈价','杠杆','TP1%','15mKOL','1HKOL','日KOL','对齐']
        n = len(sorted_results)
        table = doc.add_table(rows=1 + n, cols=len(headers))
        table.style = 'Table Grid'
        add_heading_row(table, headers)
        for i, r in enumerate(sorted_results):
            is_long = r.get('direction') == 'long'
            d_arrow = '🟢多' if is_long else '🔴空'
            hrs = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
            pos = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
            target = r['r1'] if is_long else r['s2']
            vals = [i+1, r['base'], d_arrow, f'{r["adx"]:.1f}', r.get('adx_trend',''),
                    hrs, pos, f'{r["rsi"]:.0f}',
                    fmt_price(r['entry']), fmt_price(target),
                    f'{r["okx_lev"]}x', f'{r["tp1_profit"]:.0f}%',
                    f'{r["m5_long"]}/{r["m5_short"]}',
                    f'{r.get("mtf_long",0)}/{r.get("mtf_short",0)}',
                    f'{r["daily_long"]}/{r["daily_short"]}',
                    r.get('alignment_grade', '')[:6]]
            add_data_row(table, i+1, vals)
        doc.add_paragraph()

    # ── 图例（列注释） ──
    legend_items = [
        ('序', '排名序号'),
        ('币种', '交易对'),
        ('强度', 'ADX趋势强度，≥25为强趋势'),
        ('方向', '↑增强/↓减弱'),
        ('小时', '预计ADX跌破25的小时数'),
        ('位置', 'Keltner通道位置：🔴上轨上/🟡偏高/🟢中位/🔵偏低/🟣下轨下'),
        ('RSI', '相对强弱指数，>70超买，<30超卖'),
        ('入场价', '当前价格，信号确认时参考'),
        ('止盈价', 'TP1目标价：做多=R1阻力位，做空=S2支撑位'),
        ('杠杆', 'OKX该币种最大允许杠杆'),
        ('TP1利润%', '到TP1的利润=涨幅%×杠杆，≥100%即翻倍'),
        ('15mKOL', '15分钟K线KOL投票：看多人数/看空人数'),
        ('日KOL', '日线KOL投票：看多人数/看空人数'),
    ]

    # ── 做多表格 ──
    if long_ok:
        doc.add_heading('🟢 做多推荐', level=1)
        headers = ['序','币种','强度','方向','小时','位置','RSI','入场1','加仓','止损','止盈','杠杆','TP1利润%','15mKOL','日KOL']
        table = doc.add_table(rows=1 + len(long_ok[:5]), cols=len(headers))
        table.style = 'Table Grid'
        add_heading_row(table, headers)
        for i, r in enumerate(long_ok[:5]):
            hrs = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
            pos = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
            liq = r['entry'] - (r['entry'] / r['okx_lev'])
            stop = liq + r['entry'] * 0.002
            orb_ref = max(r.get('orb_low', r['entry'] * 0.99), stop + r['entry'] * 0.001)
            vals = [i+1, r['base'], f'{r["adx"]:.1f}', r.get('adx_trend',''), hrs, pos,
                    f'{r["rsi"]:.0f}', fmt_price(r['entry']), fmt_price(orb_ref),
                    fmt_price(stop), fmt_price(r['r1']),
                    f'{r["okx_lev"]}x', f'{r["tp1_profit"]:.0f}%',
                    f'{r["m5_long"]}/{r["m5_short"]}', f'{r["daily_long"]}/{r["daily_short"]}']
            add_data_row(table, i+1, vals)
        doc.add_paragraph()

    # ── 做空表格 ──
    if short_ok:
        doc.add_heading('🔴 做空推荐', level=1)
        headers = ['序','币种','强度','方向','小时','位置','RSI','入场1','加仓','止损','止盈','杠杆','TP1利润%','15mKOL','日KOL']
        table = doc.add_table(rows=1 + len(short_ok[:5]), cols=len(headers))
        table.style = 'Table Grid'
        add_heading_row(table, headers)
        for i, r in enumerate(short_ok[:5]):
            hrs = f'{r["adx_hours_left"]:.0f}h' if r.get("adx_hours_left", 0) > 0 else ''
            pos = fmt_pos(r.get('kc_pos', 0.5), r.get('pos_high_count', 0), r.get('pos_low_count', 0), r.get('pos_lean', 0))
            liq = r['entry'] + (r['entry'] / r['okx_lev'])
            stop = liq - r['entry'] * 0.002
            orb_ref = min(r.get('orb_high', r['entry'] * 1.01), stop - r['entry'] * 0.001)
            vals = [i+1, r['base'], f'{r["adx"]:.1f}', r.get('adx_trend',''), hrs, pos,
                    f'{r["rsi"]:.0f}', fmt_price(r['entry']), fmt_price(orb_ref),
                    fmt_price(stop), fmt_price(r['s2']),
                    f'{r["okx_lev"]}x', f'{r["tp1_profit"]:.0f}%',
                    f'{r["m5_long"]}/{r["m5_short"]}', f'{r["daily_long"]}/{r["daily_short"]}']
            add_data_row(table, i+1, vals)
        doc.add_paragraph()

    # ── 分批入场参考 ──
    if passed:
        doc.add_heading('分批入场参考', level=1)
        liq_info_added = False
        for i, r in enumerate(passed[:3]):
            is_long = r['direction'] == 'long'
            dir_cn = '做多' if is_long else '做空'
            target = r['r1'] if is_long else r['s2']
            liq = r['entry'] - (r['entry'] / r['okx_lev']) if is_long else r['entry'] + (r['entry'] / r['okx_lev'])
            stop = liq + r['entry'] * 0.002 * (1 if is_long else -1)
            orb_ref = r.get('orb_low', r['entry'] * 0.99) if is_long else r.get('orb_high', r['entry'] * 1.01)
            # 入场2不能超过止损
            if is_long:
                orb_ref = max(orb_ref, stop + r['entry'] * 0.001)
            else:
                orb_ref = min(orb_ref, stop - r['entry'] * 0.001)
            p = doc.add_paragraph()
            p.add_run(f'{i+1}. {r["base"]} {dir_cn}').bold = True
            p.add_run(f'  |  入场1 {fmt_price(r["entry"])}(20U)')
            p.add_run(f'  →  入场2 {fmt_price(orb_ref)}(30U) ORB{"下沿" if is_long else "上沿"}')
            p.add_run(f'  |  止损 {fmt_price(stop)}')
            p.add_run(f'  |  TP1 {fmt_price(target)}')
            hrs = r.get('adx_hours_left', 0)
            if hrs > 0:
                p.add_run(f'  |  趋势预估 {hrs:.0f}h')

    # ── 列注释 ──
    doc.add_paragraph()
    doc.add_heading('列注释说明', level=2)
    for name, desc in legend_items:
        p = doc.add_paragraph()
        p.add_run(f'  {name}：').bold = True
        p.add_run(desc)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    # 保存
    docx_path = os.path.join(QF, f'daily_picks_{datetime.now().strftime("%Y-%m-%d_%H%M")}.docx')
    doc.save(docx_path)
    print(f'  已保存: {docx_path}')


if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument('--top', type=int, default=50)
    p.add_argument('--min-r1', type=float, default=1.5)
    p.add_argument('--min-oi', type=float, default=600000)
    p.add_argument('--coins', type=str, help='指定币种(逗号分隔)，如 BCH,LPT,UNI')
    a = p.parse_args()
    coins_list = a.coins.split(',') if a.coins else None
    run(top_n=a.top, min_r1=a.min_r1, min_oi=int(a.min_oi), coins=coins_list)
